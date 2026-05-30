"""
Formal Dispute Letter Generator for BESCOM/HESCOM
Produces a properly formatted DOCX letter with all required sections
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import io


def _set_cell_bg(cell, hex_color: str):
    """Set background color of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def generate_dispute_letter(
    consumer_name: str,
    consumer_number: str,
    consumer_address: str,
    discom: str,
    anomalous_months: list,  # list of dicts with month, units, amount, cause
    stats: dict,
    ai_explanation: str,
    output_path: str = None,
) -> bytes:
    """
    Generate a formal dispute letter as DOCX.
    Returns bytes so it can be downloaded from Streamlit.
    """
    doc = Document()

    # ── Page margins ────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    # ── Styles ──────────────────────────────────────────────────────────────
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)

    def add_heading(text, level=1, color="1F3864"):
        para = doc.add_heading(text, level=level)
        run = para.runs[0] if para.runs else para.add_run(text)
        run.font.color.rgb = RGBColor.from_string(color)
        run.font.bold = True
        return para

    def add_para(text, bold=False, italic=False, size=11, align=None):
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.name = "Calibri"
        if align:
            para.alignment = align
        return para

    def add_field(label, value):
        para = doc.add_paragraph()
        run_label = para.add_run(f"{label}: ")
        run_label.bold = True
        run_label.font.name = "Calibri"
        run_label.font.size = Pt(11)
        run_val = para.add_run(str(value))
        run_val.font.name = "Calibri"
        run_val.font.size = Pt(11)
        return para

    # ── Header ──────────────────────────────────────────────────────────────
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run(f"COMPLAINT / DISPUTE LETTER")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor.from_string("1F3864")

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run("Regarding Inflated / Erroneous Electricity Bill")
    sub_run.font.size = Pt(12)
    sub_run.italic = True
    sub_run.font.name = "Calibri"
    sub_run.font.color.rgb = RGBColor.from_string("C00000")

    doc.add_paragraph()

    # ── Date & Reference ────────────────────────────────────────────────────
    today = datetime.date.today().strftime("%d %B %Y")
    add_para(f"Date: {today}", bold=False)
    add_para(f"Reference No.: DISPUTE/{consumer_number}/{datetime.date.today().strftime('%Y%m%d')}")
    doc.add_paragraph()

    # ── Addressee ───────────────────────────────────────────────────────────
    discom_addresses = {
        "BESCOM": (
            "The Executive Engineer / Consumer Grievance Cell\n"
            "Bangalore Electricity Supply Company Limited (BESCOM)\n"
            "K.R. Circle, Bangalore – 560 001\n"
            "Karnataka, India\n"
            "Email: ccc@bescom.co.in | Helpline: 1912"
        ),
        "HESCOM": (
            "The Executive Engineer / Consumer Grievance Cell\n"
            "Hubli Electricity Supply Company Limited (HESCOM)\n"
            "Navanagar, Hubli – 580 025\n"
            "Karnataka, India\n"
            "Email: hescomhelpline@hescom.co.in | Helpline: 1912"
        ),
    }
    addressee_text = discom_addresses.get(discom, discom_addresses["BESCOM"])
    add_para("To,", bold=True)
    add_para(addressee_text)
    doc.add_paragraph()

    # ── Subject Line ────────────────────────────────────────────────────────
    anomaly_months_str = ", ".join([a["month"] for a in anomalous_months])
    subject_para = doc.add_paragraph()
    subject_run = subject_para.add_run(
        f"Subject: Formal Complaint Against Inflated Electricity Bill for "
        f"{anomaly_months_str} — Consumer No. {consumer_number}"
    )
    subject_run.bold = True
    subject_run.font.size = Pt(11)
    subject_run.font.name = "Calibri"
    subject_run.font.color.rgb = RGBColor.from_string("1F3864")

    doc.add_paragraph()

    # ── Salutation ──────────────────────────────────────────────────────────
    add_para("Respected Sir/Madam,")
    doc.add_paragraph()

    # ── Section 1: Consumer Details ─────────────────────────────────────────
    add_heading("1. Consumer Details", level=2)
    add_field("Consumer Name", consumer_name)
    add_field("Consumer / Account Number", consumer_number)
    add_field("Service Address", consumer_address)
    add_field("DISCOM / Utility", discom)
    doc.add_paragraph()

    # ── Section 2: Nature of Complaint ──────────────────────────────────────
    add_heading("2. Nature of Complaint", level=2)
    add_para(
        "I, the undersigned consumer, am writing to formally dispute the electricity bill(s) "
        f"issued for the month(s) of {anomaly_months_str}. The billed consumption and corresponding "
        "charges are significantly higher than my 12-month historical average and cannot be "
        "explained by genuine consumption changes. The anomalies were identified through "
        "statistical analysis (Z-Score and IQR method) of my billing history."
    )
    doc.add_paragraph()

    # ── Section 3: Statistical Evidence ─────────────────────────────────────
    add_heading("3. Statistical Evidence of Anomaly", level=2)
    add_para(
        f"The following statistics from my 12-month billing history (AI-assisted analysis) "
        "confirm the disputed month(s) are statistical outliers:"
    )
    doc.add_paragraph()

    stats_table = doc.add_table(rows=1, cols=2)
    stats_table.style = "Table Grid"
    hdr_cells = stats_table.rows[0].cells
    hdr_cells[0].text = "Metric"
    hdr_cells[1].text = "Value"
    for cell in hdr_cells:
        _set_cell_bg(cell, "1F3864")
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor.from_string("FFFFFF")
        run.font.size = Pt(10)

    stat_rows = [
        ("12-Month Average Consumption", f"{stats.get('mean_units', 'N/A')} units"),
        ("Standard Deviation", f"{stats.get('std_units', 'N/A')} units"),
        ("IQR Normal Range (Lower–Upper)", f"{stats.get('iqr_lower', 'N/A')} – {stats.get('iqr_upper', 'N/A')} units"),
        ("Minimum / Maximum (12 months)", f"{stats.get('min_units', 'N/A')} / {stats.get('max_units', 'N/A')} units"),
    ]
    for label, value in stat_rows:
        row = stats_table.add_row().cells
        row[0].text = label
        row[1].text = value
        for cell in row:
            cell.paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # ── Section 4: Disputed Bills Table ─────────────────────────────────────
    add_heading("4. Disputed Bill Details", level=2)

    bill_table = doc.add_table(rows=1, cols=5)
    bill_table.style = "Table Grid"
    headers = ["Month", "Units Billed", "Amount Billed (₹)", "Expected Range", "Likely Cause"]
    for i, h in enumerate(headers):
        cell = bill_table.rows[0].cells[i]
        cell.text = h
        _set_cell_bg(cell, "C00000")
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor.from_string("FFFFFF")
        run.font.size = Pt(9)

    for anomaly in anomalous_months:
        row = bill_table.add_row().cells
        row[0].text = anomaly.get("month", "")
        row[1].text = str(anomaly.get("units", ""))
        row[2].text = f"₹ {anomaly.get('amount', '')}"
        row[3].text = f"{stats.get('iqr_lower', '')}–{stats.get('iqr_upper', '')} units"
        row[4].text = anomaly.get("cause", "")
        for cell in row:
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph()

    # ── Section 5: AI-Assisted Explanation ──────────────────────────────────
    add_heading("5. Technical Explanation & Root Cause Analysis", level=2)
    add_para(
        "An AI-assisted analysis of my 12-month billing data identified the following "
        "potential causes for the anomalous billing:"
    )
    doc.add_paragraph()

    # Split the AI explanation into paragraphs
    for para_text in ai_explanation.split("\n"):
        if para_text.strip():
            add_para(para_text.strip())

    doc.add_paragraph()

    # ── Section 6: Relief Sought ─────────────────────────────────────────────
    add_heading("6. Relief / Action Requested", level=2)
    reliefs = [
        "Immediate review and correction of the disputed bill(s) based on actual/verified meter readings.",
        "Physical inspection and testing of the electricity meter at my premises within 30 days as per the Electricity Act, 2003 (Section 26).",
        "Provisional bill based on my 12-month average consumption until the investigation is complete.",
        "Refund or adjustment of any excess amount charged in the disputed month(s).",
        "Written explanation of the cause of excess billing within 15 working days.",
        "Ensure no disconnection of supply during the pendency of this dispute (per KERC Regulation 5(5)).",
    ]
    for i, relief in enumerate(reliefs, 1):
        para = doc.add_paragraph(style="List Number")
        run = para.add_run(relief)
        run.font.size = Pt(11)
        run.font.name = "Calibri"

    doc.add_paragraph()

    # ── Section 7: Legal Reference ───────────────────────────────────────────
    add_heading("7. Applicable Legal Provisions", level=2)
    legal_refs = [
        "Electricity Act, 2003 — Section 26 (Meter testing & inspection)",
        "Electricity Act, 2003 — Section 135 (Theft provisions, inapplicable here)",
        "KERC (Consumer Grievance Redressal Forum & Ombudsman) Regulations, 2004",
        "Consumer Protection Act, 2019 — Section 2(1)(d) deficiency in service",
        "CERC/SERC Billing Code — Right to accurate billing",
    ]
    for ref in legal_refs:
        para = doc.add_paragraph(style="List Bullet")
        run = para.add_run(ref)
        run.font.size = Pt(10)
        run.font.name = "Calibri"

    doc.add_paragraph()

    # ── Section 8: Documents Enclosed ───────────────────────────────────────
    add_heading("8. Enclosures", level=2)
    enclosures = [
        "Copy of disputed electricity bill(s)",
        "12-month billing history (CSV data printout)",
        "Statistical anomaly analysis report (AI-generated)",
        "Previous bills for reference (last 6 months)",
        "Consumer identity proof (Aadhaar / Voter ID)",
    ]
    for i, enc in enumerate(enclosures, 1):
        para = doc.add_paragraph(style="List Number")
        run = para.add_run(enc)
        run.font.size = Pt(10)
        run.font.name = "Calibri"

    doc.add_paragraph()

    # ── Closing ──────────────────────────────────────────────────────────────
    add_para(
        "I trust that this matter will be treated with urgency and resolved amicably. "
        "I am available for any on-site inspection or clarification at mutually convenient time. "
        "If the matter is not resolved within 30 days, I reserve the right to approach the "
        "Karnataka Electricity Regulatory Commission (KERC) / Consumer Forum."
    )
    doc.add_paragraph()
    add_para("Thanking you,")
    doc.add_paragraph()
    doc.add_paragraph()
    add_para("Yours faithfully,")
    doc.add_paragraph()
    doc.add_paragraph()
    add_para("_______________________________", bold=True)
    add_field("Name", consumer_name)
    add_field("Consumer No.", consumer_number)
    add_field("Address", consumer_address)
    add_field("Date", today)
    add_para("(Signature)")

    # ── Footer note ──────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        "This letter was prepared with AI-assisted bill anomaly analysis. "
        "Complaints can also be filed online at: grievance.karnataka.gov.in | KERC: kerc.kar.nic.in"
    )
    footer_run.font.size = Pt(8)
    footer_run.italic = True
    footer_run.font.color.rgb = RGBColor.from_string("666666")
    footer_run.font.name = "Calibri"

    # ── Save ─────────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    if output_path:
        doc.save(output_path)
    return buf.getvalue()
